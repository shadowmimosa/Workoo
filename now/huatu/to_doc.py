import re
import os
import json
import requests
from bson import ObjectId
from urllib import parse
from docx import Document
from docx.shared import Inches

from config import MONGO
from utils.log import logger
from utils.run import run_func
from pymongo import MongoClient


class MongoOpea(object):
    def __init__(self):
        self.init_mongo()
        super().__init__()

    def init_mongo(self):

        config = MONGO["debug"]

        config["user"] = parse.quote_plus(config["user"])
        config["passwd"] = parse.quote_plus(config["passwd"])

        client = MongoClient(
            "mongodb://{user}:{passwd}@{host}:{port}/".format(**config),
            connect=False)

        self.mongo = client["huatu"]['raw_data']

    def select(self, point):
        # return list(
        #     self.mongo.find({'_id': ObjectId('5e694dc6d5e714ac48e3d8a2')}))
        return list(self.mongo.find({"pointList.points": point}))

    def update(self, sync):
        pass

    def insert(self, data):
        pass


def get_data(point):
    results = mongo.select(point)

    for data in results:
        yield data


def down_load_img(url, name):
    r = requests.get(url)
    r.raise_for_status()
    with open(name, 'wb') as fp:
        fp.write(r.content)
        fp.close()


def replace(content):
    return content.replace('<p>', '').replace('</p>', '').replace(
        '<br>', '\n').replace('<br/>', '\n')


def img_handle(content):
    if '<img' in content:

        imgs = re.findall('<img src="(.*?)" width="[0-9]*" height="[0-9]*">',
                          content)
        content = re.sub('<img src="(.*?)" width="[0-9]*" height="[0-9]*">',
                         '', content)
        content = re.sub('\（?\(?[\\xa0]*\（?\(?', '', content)


def content_handle(content, index):
    if not content:
        return

    if '<img=' in content and '></img>' in content:
        imgs_url = re.findall('<img=(.*?)></img>', content)
        for url in imgs_url:

            question = content.replace('\\xa0', '').replace('\\n', '').replace(
                '<img=' + url + '></img>', '@@@')

            document.add_paragraph(replace(f'{index}. {question}'))

        for url in imgs_url:

            img_name = f'./data/pic/{url.split("/")[-1]}'
            down_load_img(url, img_name)
            document.add_picture(img_name, width=Inches(2.5))
    else:
        question = content.replace('\\xa0', '').replace('\\n', '')
        document.add_paragraph(replace(f'{index}. {question}'))


def material_handle(content):
    if not content:
        return

    if '<img=' in content and '></img>' in content:
        imgs_url = re.findall('<img=(.*?)></img>', content)
        for url in imgs_url:

            question = content.replace('\\xa0', '').replace('\\n', '').replace(
                '<img=' + url + '></img>', '@@@')

            document.add_paragraph(replace(question))

        for url in imgs_url:

            img_name = f'./data/pic/{url.split("/")[-1]}'
            down_load_img(url, img_name)
            document.add_picture(img_name, width=Inches(2.5))
    else:
        question = content.replace('\\xa0', '').replace('\\n', '')
        document.add_paragraph(replace(question))


def option_handle(content, option):
    if not content:
        return

    if '<img=' in content and '></img>' in content:
        imgs_url = re.findall('<img=(.*?)></img>', content)
        for url in imgs_url:

            analysis = content.replace('\\xa0', '').replace('\\n', '').replace(
                '<img=' + url + '></img>', '@@@')

            document.add_paragraph(replace(f'解析: {analysis}'))

        for url in imgs_url:

            img_name = f'./data/pic/{url.split("/")[-1]}'
            down_load_img(url, img_name)
            document.add_picture(img_name, width=Inches(1))
    else:
        question = content.replace('\\xa0', '').replace('\\n', '')
        document.add_paragraph(replace(f'{option}: {question}'))


def analysis_handle(content):
    if not content:
        return

    if '<img=' in content and '></img>' in content:
        imgs_url = re.findall('<img=(.*?)></img>', content)
        for url in imgs_url:

            analysis = content.replace('\\xa0', '').replace('\\n', '').replace(
                '<img=' + url + '></img>', '@@@')

            document.add_paragraph(replace(f'解析: {analysis}'))

        for url in imgs_url:

            img_name = f'./data/pic/{url.split("/")[-1]}'
            down_load_img(url, img_name)
            document.add_picture(img_name, width=Inches(1))
    else:
        analysis = content.replace('\\xa0', '').replace('\\n', '')
        document.add_paragraph(replace(f'解析: {analysis}'))


def end_handle(points, source):
    if points:
        document.add_paragraph(replace(f'考点: {" ".join(points)}'))
    if source:
        document.add_paragraph(replace(f'来源: {source}'))

    document.add_paragraph('\n\n\n')


def kind_handle(data):
    kind = data.get('type')
    choices = data.get('choices')
    answer = data.get('answer')
    percents = data.get('meta').get('percents')
    if len(percents) == 1:
        percents = percents[0]
    else:
        percents = percents[data.get('meta').get('answers').index(answer)]
    if not choices:
        return
    # "99": "单选题",
    # "100": "多选题",
    # "101": "不定项选择",
    # "105": "复合题",
    # "109": "判断题"
    if kind == 99:
        option_handle(choices[0], 'A')
        option_handle(choices[1], 'B')
        option_handle(choices[2], 'C')
        option_handle(choices[3], 'D')
        document.add_paragraph(
            f'答案: {str(answer).replace("1","A ").replace("2","B ").replace("3","C ").replace("4","D ")}'
        )
        document.add_paragraph(f'全站准确率: {percents}')
    elif kind == 100:
        option_handle(choices[0], 'A')
        option_handle(choices[1], 'B')
        option_handle(choices[2], 'C')
        option_handle(choices[3], 'D')
        answer = str(answer)
        temp = 0
        # for index in range(1, 6):
        #     if str(index) in answer:
        #         temp += percents[index - 1]

        answer = f'答案: {answer.replace("1","A ").replace("2","B ").replace("3","C ").replace("4","D ")}'
        document.add_paragraph(answer)
        document.add_paragraph(f'全站准确率: {percents}')
    elif kind == 101:
        option_handle(choices[0], 'A')
        option_handle(choices[1], 'B')
        option_handle(choices[2], 'C')
        option_handle(choices[3], 'D')
        answer = str(answer)
        temp = 0
        # for index in range(1, 6):
        #     if str(index) in answer:
        #         temp += percents[index - 1]

        answer = f'答案: {answer.replace("1","A ").replace("2","B ").replace("3","C ").replace("4","D ")}'
        document.add_paragraph(answer)
        document.add_paragraph(f'全站准确率: {percents}')
    elif kind == 105:
        option_handle(choices[0], 'A')
        option_handle(choices[1], 'B')
        option_handle(choices[2], 'C')
        option_handle(choices[3], 'D')
        answer = str(answer)
        temp = 0
        # for index in range(1, 6):
        #     if str(index) in answer:
        #         temp += percents[index - 1]

        answer = f'答案: {answer.replace("1","A ").replace("2","B ").replace("3","C ").replace("4","D ")}'
        document.add_paragraph(answer)
        document.add_paragraph(f'全站准确率: {percents}')
    elif kind == 109:
        option_handle(choices[0], 'A')
        option_handle(choices[1], 'B')
        document.add_paragraph(
            f'答案: {str(answer).replace("1","A ").replace("2","B ").replace("3","C ").replace("4","D ")}'
        )
        document.add_paragraph(f'全站准确率: {percents}')
    else:
        logger.error(f'No type catch, {kind}')


def main_handle(data, index):
    if not data.get('pointsName'):
        logger.error(f'{data} is wrong')
        return

    content_handle(data.get('stem'), index)
    if data.get('materials'):
        for item in data.get('materials'):
            material_handle(item)

    kind_handle(data)

    analysis_handle(data.get('analysis'))
    end_handle(data.get('pointsName'), data.get('from'))


class GetPoint(object):
    def __init__(self, content):
        self.need = []
        self.get_all(content)
        super().__init__()

    def get_all(self, content):
        if content.get('child') is None:
            for key in content:
                self.get_all(content[key])

        elif not content.get('child'):
            self.need.append(content.get('point'))
        else:
            self.get_all(content.get('child'))


def main(ids):
    global document
    img_handle(
        '如下图所示，完全相同的两根弹簧，下面挂两个质量相同、形状不同的实心铁块，其中甲是立方体，乙是球体。现将两个铁块完全浸没在某盐水溶液中，该溶液的密度随深度增加而均匀增加。待两铁块静止后，甲、乙两铁块受到的弹簧的拉力相比（  ）。<img src="http://tiku.huatu.com/cdn/pandora/img/5ac36cb6-b1b7-4da9-98e8-bc21269bdf24..png?imageView2/0/w/643/format/jpg" width="642" height="417">'
    )
    # with open('./data/points.json', 'r', encoding='utf-8') as fn:
    #     data = json.loads(fn.read())

    # kindname = '事业单位考试'
    kindname = '公务员行测'
    # data.pop('公务员行测')
    # data = data.pop('公务员行测')
    # points = set(GetPoint(data).need)

    # with open('./point.json', 'w', encoding='utf-8') as fn:
    #     fn.write(json.dumps(list(points)))

    count = 0
    for point in ["674", "65695", "65748", "433", "431", "1006", "65763"]:
        # for point in points:
        document = Document()

        # yield_data = get_data(int(65879))
        yield_data = get_data(int(point))

        try:
            data = next(yield_data)
        except StopIteration:
            continue

        pointsname = data.get("pointsName")
        if not pointsname:
            logger.warning(f'pointList is null - {data.get("_id")}')
            continue
            data['pointsName'] = data.get('pointList')[0].get('pointsName')
            pointsname = data['pointsName']

        dirname = f'./data/doc/{kindname}/{"/".join(pointsname[:-1])}/'
        document.add_heading(data.get('pointsName')[-1])
        try:
            os.makedirs(dirname)
        except FileExistsError:
            pass
        index = 1

        while True:
            # main_handle(data, index)
            # try:
            #     data = next(yield_data)
            # except StopIteration:
            #     break
            try:
                main_handle(data, index)
            except Exception as exc:
                logger.error(
                    f'{data.get("_id")}is error in main handle, {exc}')
            finally:
                index += 1
                count += 1
                try:
                    data = next(yield_data)
                except StopIteration:
                    yield_data.close()
                    break
            logger.info(f'{index} is down')

        print(f'counts is {count}')
        document.save(f'{dirname}{pointsname[-1]}.docx')


if __name__ == "__main__":
    mongo = MongoOpea()
    document = Document()
    main(798)