import os
import re
import requests
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError, UnexpectedEndOfFileError


def down_img(url):
    filename = f'./data/pic/{url.split("?")[0].split("/")[-1].replace("..",".")}'
    if os.path.exists(filename):
        return filename

    resp = session.get(url)
    if resp.status_code != 200:
        print(resp)
        return

    with open(filename, 'wb') as fn:
        fn.write(resp.content)

    return filename


def img_handle(content):
    # content = re.sub('\（?\(?[\\xa0]*\（?\(?', '', content)
    content = re.sub('[\\xa0]*', '', content)
    img_url = []
    paths = []

    if '<img' in content:
        imgs = re.findall('<img[\s\S]*src="(.*?)"[\s\S]*>', content)
        content = re.sub('<img[\s\S]*src="(.*?)"[\s\S]*>', '', content)
        for img in imgs:
            try:
                result = down_img(img)
            except Exception as exc:
                print(exc)
            else:
                paths.append(result)

    return content, paths


def to_doc(document: Document, contents):
    for index, content in enumerate(contents):
        if index == 0:
            document.add_heading(content)
        else:
            content, imgs = img_handle(content)
            document.add_paragraph(content)
            for img in imgs:
                try:
                    document.add_picture(img, width=Inches(2.5))
                except UnrecognizedImageError:
                    print('image error')
                except ZeroDivisionError:
                    print('zero error')
                except UnexpectedEndOfFileError:
                    print('end error')
                except Exception as exc:
                    print('error is {}'.format(exc))


def count(lines: list, dirpath):
    header = lines[0]
    lines.reverse()
    for index, line in enumerate(lines):
        if line == '\n\n\n' and index != 0:
            with open('./counts.txt', 'a', encoding='utf-8') as fn:
                fn.write(
                    f'{dirpath} - {header} - {lines[index-1].split(". ")[0]}')
                fn.write('\n')
                return


def main():
    for root, dirs, files in os.walk('./data/doc'):
        for filename in files:
            # document = Document(
            #     r'C:\Users\ShadowMimosa\Documents\Repos\Workoo\now\huatu\data\doc\公务员行测\判断推理\科学推理\科学推理物理.docx'
            # )
            dirpath = root.replace('doc', 'doc_new').replace('\\', '/')
            filepath = os.path.join(dirpath, filename)
            if os.path.isfile(filepath):
                continue
            document = Document(os.path.join(root, filename))
            lines = [paragraph.text for paragraph in document.paragraphs]
            new_document = Document()
            to_doc(new_document, lines)
            try:
                os.makedirs(dirpath)
            except FileExistsError:
                pass
            new_document.save(filepath)
            count(lines, dirpath)


if __name__ == "__main__":
    session = requests.session()
    session.headers = {
        'User-Agent':
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.100 Safari/537.36'
    }
    session.verify = False
    main()
